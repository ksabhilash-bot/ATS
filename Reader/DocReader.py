from pathlib import Path

from docx import Document
from pypdf import PdfReader

from utils.text_cleaner import clean_resume_text


SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


class ResumeExtractionError(ValueError):
    """Raised when an unsupported resume format is provided."""
    pass


def extract_resume_text(file: str | Path) -> str:
    """
    Extract text from a resume file based on its extension.

    Args:
        file: Path to the resume file.

    Returns:
        Extracted and cleaned resume text.

    Raises:
        ResumeExtractionError: If the file type is unsupported.
    """
    path = Path(file)

    if not path.exists():
        raise FileNotFoundError(f"Resume file not found: {path}")

    extension = path.suffix.lower()

    if extension == ".pdf":
        return extract_pdf_text(path)

    if extension == ".docx":
        return extract_docx_text(path)

    supported = ", ".join(sorted(SUPPORTED_EXTENSIONS))

    raise ResumeExtractionError(
        f"Unsupported resume file type '{extension or 'no extension'}'. "
        f"Supported types: {supported}"
    )


def extract_pdf_text(file: str | Path) -> str:
    """
    Extract text from a PDF resume.
    """
    try:
        reader = PdfReader(str(file))

        page_text = []

        for page in reader.pages:
            page_text.append(page.extract_text() or "")

        text = "\n".join(
            text for text in page_text if text
        ).strip()

        return clean_resume_text(text)

    except Exception as error:
        raise ResumeExtractionError(
            f"Failed to extract text from PDF: {error}"
        ) from error


def extract_docx_text(file: str | Path) -> str:
    """
    Extract text from a DOCX resume, including tables.
    """
    try:
        document = Document(str(file))

        blocks = []

        # Extract normal paragraphs
        blocks.extend(
            _paragraph_text(paragraph)
            for paragraph in document.paragraphs
        )

        # Extract tables
        blocks.extend(
            _table_text(table)
            for table in document.tables
        )

        text = "\n".join(
            block for block in blocks if block
        ).strip()

        return clean_resume_text(text)

    except Exception as error:
        raise ResumeExtractionError(
            f"Failed to extract text from DOCX: {error}"
        ) from error


def _paragraph_text(paragraph) -> str:
    """
    Extract text from a paragraph.
    """
    return paragraph.text.strip()


def _table_text(table) -> str:
    """
    Extract text from a table.

    Output format:

    Company | Role
    Acme     | Engineer
    """
    rows = []

    for row in table.rows:
        cells = [_cell_text(cell) for cell in row.cells]

        cells = [cell for cell in cells if cell]

        if cells:
            rows.append(" | ".join(cells))

    return "\n".join(rows)


def _cell_text(cell) -> str:
    """
    Extract text from a table cell.
    """
    parts = []

    # Paragraphs inside the cell
    parts.extend(
        _paragraph_text(paragraph)
        for paragraph in cell.paragraphs
    )

    # Nested tables inside the cell
    parts.extend(
        _table_text(table)
        for table in cell.tables
    )

    return " ".join(
        part for part in parts if part
    ).strip()