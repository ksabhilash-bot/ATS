from pathlib import Path

from docx import Document
import pytest

from Reader.DocReader import (
    ResumeExtractionError,
    extract_docx_text,
    extract_resume_text,
)


def _write_docx(path: Path, body_xml: str) -> None:
    document = Document()
    for line in body_xml.splitlines():
        text = line.strip()
        if text:
            document.add_paragraph(text)
    document.save(path)


def _write_docx_with_table(path: Path) -> None:
    document = Document()
    document.add_paragraph("Experience")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Company"
    table.cell(0, 1).text = "Role"
    table.cell(1, 0).text = "Acme"
    table.cell(1, 1).text = "Engineer"

    document.save(path)


def test_extract_resume_text_routes_docx_by_extension(tmp_path):
    resume = tmp_path / "resume.docx"
    _write_docx(
        resume,
        """
        Jane Candidate
        Python Developer
        """,
    )

    assert extract_resume_text(resume) == "Jane Candidate\nPython Developer"


def test_extract_docx_text_includes_table_rows_and_cells(tmp_path):
    resume = tmp_path / "resume.docx"
    _write_docx_with_table(resume)

    assert extract_docx_text(resume) == (
        "Experience\nCompany | Role\nAcme | Engineer"
    )


def test_extract_resume_text_rejects_unsupported_extension(tmp_path):
    resume = tmp_path / "resume.txt"
    resume.write_text("plain text", encoding="utf-8")

    with pytest.raises(ResumeExtractionError, match="Unsupported resume file type"):
        extract_resume_text(resume)
